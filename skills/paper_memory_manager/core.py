"""
Paper Memory Manager - Core Modules

包含三个核心模块：
- HeadingDetector: 章节识别器（混合策略）
- SectionExtractor: 内容提取器（层级包含）
- MemoryUpdater: 记忆更新器（diff生成+审核）
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from docx import Document


@dataclass
class Heading:
    """章节标题数据类"""
    level: int                    # 层级 1-4
    title: str                    # 标题文本
    numbering: List[int]          # 编号数组 [2,1] 表示 2.1
    numbering_str: str            # 编号字符串 "2.1"
    start_para: int               # 起始段落索引
    end_para: int = 0             # 结束段落索引（由后续章节决定）
    confidence: float = 1.0       # 识别置信度

    def __post_init__(self):
        if not self.numbering_str:
            self.numbering_str = '.'.join(map(str, self.numbering))


@dataclass
class HeadingTree:
    """章节树"""
    headings: List[Heading] = field(default_factory=list)
    doc_path: str = ""

    def find_section(self, section_id: str) -> Optional[Heading]:
        """查找指定章节"""
        # 支持 "2.1" 或 "第2章" 格式
        if section_id.startswith('第') and '章' in section_id:
            # 解析 "第2章"
            match = re.search(r'第(\d+)章', section_id)
            if match:
                chapter_num = int(match.group(1))
                for h in self.headings:
                    if h.level == 1 and h.numbering == [chapter_num]:
                        return h
        else:
            # 解析 "2.1"、"2.1.3"
            parts = section_id.split('.')
            target = [int(p) for p in parts]
            for h in self.headings:
                if h.numbering == target:
                    return h
        return None

    def get_level1_sections(self) -> List[Heading]:
        """获取所有一级章节"""
        return [h for h in self.headings if h.level == 1]

    def get_children(self, heading: Heading) -> List[Heading]:
        """获取指定章节的所有子节"""
        children = []
        found = False
        for h in self.headings:
            if h == heading:
                found = True
                continue
            if found:
                # 判断是否为目标章节的子节
                if len(h.numbering) > len(heading.numbering):
                    if h.numbering[:len(heading.numbering)] == heading.numbering:
                        children.append(h)
                else:
                    # 同级或更高级别，停止
                    break
        return children


class HeadingDetector:
    """
    章节识别器 - 混合策略

    1. Word样式优先（Heading 1-4）
    2. 文本模式回退（正则匹配）
    3. 宽松容错修复
    """

    # 文本模式识别规则
    # 支持编号后无空格的情况: "1.1.1标题" 或 "1.1.1 标题"
    # 注意：按层级从高到低排序，避免 "1.1.1" 被误判为 "1.1"
    CHAPTER_PATTERNS = [
        # 四级: X.X.X.X (最高优先级)
        (4, r'^(\d+)\.(\d+)\.(\d+)\.(\d+)\s*[\.\s\-:]?(.*)$'),
        # 三级: X.X.X
        (3, r'^(\d+)\.(\d+)\.(\d+)\s*[\.\s\-:]?(.*)$'),
        # 二级: X.X
        (2, r'^(\d+)\.(\d+)\s*[\.\s\-:]?(.*)$'),
        # 一级: 第X章 (最低优先级)
        (1, r'^第\s*([一二三四五六七八九十\d]+)\s*章\s*[\.\s\-:]?(.*)$'),
    ]

    # 中文数字映射
    CHINESE_NUMBERS = {
        '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
        '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
        '0': 0, '1': 1, '2': 2, '3': 3, '4': 4,
        '5': 5, '6': 6, '7': 7, '8': 8, '9': 9
    }

    def detect(self, docx_path: str) -> HeadingTree:
        """
        检测章节结构

        Args:
            docx_path: DOCX文件路径

        Returns:
            HeadingTree: 章节树
        """
        doc = Document(docx_path)

        # 第一优先：Word样式
        headings = self._detect_by_style(doc)

        # 第二优先：文本模式（样式失败时）
        if not headings:
            headings = self._detect_by_text(doc)

        # 容错修复
        headings = self._repair_numbering(headings)

        # 计算每个章节的结束位置
        headings = self._calculate_end_positions(headings, len(doc.paragraphs))

        return HeadingTree(headings=headings, doc_path=docx_path)

    def _detect_by_style(self, doc: Document) -> List[Heading]:
        """通过Word样式检测章节"""
        headings = []

        for i, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else 'Normal'
            text = para.text.strip()

            if not text:
                continue

            level = None
            numbering = []

            # 检测样式
            if 'Heading 1' in style_name or style_name == 'Heading 1':
                level = 1
                # 尝试从文本解析编号
                match = re.match(r'^第\s*(\d+)\s*章', text)
                if match:
                    numbering = [int(match.group(1))]
                else:
                    # 没有编号，按顺序分配
                    numbering = [len([h for h in headings if h.level == 1]) + 1]

            elif 'Heading 2' in style_name or style_name == 'Heading 2':
                level = 2
                match = re.match(r'^(\d+)\.(\d+)', text)
                if match:
                    numbering = [int(match.group(1)), int(match.group(2))]
                else:
                    # 尝试推断
                    numbering = self._infer_numbering(headings, 2)

            elif 'Heading 3' in style_name or style_name == 'Heading 3':
                level = 3
                match = re.match(r'^(\d+)\.(\d+)\.(\d+)', text)
                if match:
                    numbering = [int(match.group(1)), int(match.group(2)), int(match.group(3))]
                else:
                    numbering = self._infer_numbering(headings, 3)

            elif 'Heading 4' in style_name or style_name == 'Heading 4':
                level = 4
                match = re.match(r'^(\d+)\.(\d+)\.(\d+)\.(\d+)', text)
                if match:
                    numbering = [int(match.group(1)), int(match.group(2)),
                                int(match.group(3)), int(match.group(4))]
                else:
                    numbering = self._infer_numbering(headings, 4)

            if level:
                headings.append(Heading(
                    level=level,
                    title=text,
                    numbering=numbering,
                    numbering_str='.'.join(map(str, numbering)),
                    start_para=i,
                    confidence=0.9
                ))

        return headings

    def _detect_by_text(self, doc: Document) -> List[Heading]:
        """通过文本模式检测章节（样式失败时回退）"""
        headings = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            for level, pattern in self.CHAPTER_PATTERNS:
                match = re.match(pattern, text)
                if match:
                    groups = match.groups()

                    if level == 1:
                        # 第X章
                        num_str = groups[0]
                        if num_str.isdigit():
                            chapter_num = int(num_str)
                        else:
                            chapter_num = self._chinese_to_number(num_str)
                        numbering = [chapter_num]
                        title = groups[1] if len(groups) > 1 else text

                    elif level == 2:
                        numbering = [int(groups[0]), int(groups[1])]
                        title = groups[2] if len(groups) > 2 else text

                    elif level == 3:
                        numbering = [int(groups[0]), int(groups[1]), int(groups[2])]
                        title = groups[3] if len(groups) > 3 else text

                    elif level == 4:
                        numbering = [int(groups[0]), int(groups[1]),
                                    int(groups[2]), int(groups[3])]
                        title = groups[4] if len(groups) > 4 else text

                    headings.append(Heading(
                        level=level,
                        title=text,
                        numbering=numbering,
                        numbering_str='.'.join(map(str, numbering)),
                        start_para=i,
                        confidence=0.7
                    ))
                    break  # 匹配成功，不再检查其他模式

        return headings

    def _chinese_to_number(self, chinese: str) -> int:
        """中文数字转阿拉伯数字"""
        result = 0
        temp = 0
        for char in chinese:
            if char in self.CHINESE_NUMBERS:
                num = self.CHINESE_NUMBERS[char]
                if num == 10:
                    if temp == 0:
                        temp = 1
                    result += temp * 10
                    temp = 0
                else:
                    temp = temp * 10 + num if temp else num
        result += temp
        return result if result > 0 else 1

    def _infer_numbering(self, headings: List[Heading], level: int) -> List[int]:
        """推断章节编号"""
        # 获取同级别的上一个章节
        same_level = [h for h in headings if h.level == level]

        if same_level:
            # 递增编号
            last = same_level[-1]
            new_numbering = last.numbering.copy()
            new_numbering[-1] += 1
            return new_numbering
        else:
            # 第一个该级别的章节
            if level == 1:
                return [1]
            else:
                # 找到父级章节
                parent_level = level - 1
                parent_sections = [h for h in headings if h.level == parent_level]
                if parent_sections:
                    parent = parent_sections[-1]
                    return parent.numbering + [1]
                return [1] * level

    def _repair_numbering(self, headings: List[Heading]) -> List[Heading]:
        """
        宽松容错修复

        - 补全缺失编号
        - 修复编号跳跃
        - 清理无效标题
        """
        if not headings:
            return headings

        # 按段落顺序排序
        headings.sort(key=lambda h: h.start_para)

        repaired = []
        prev_level = 0
        prev_numbering = []

        for h in headings:
            # 跳过明显无效的标题（太短或太长）
            if len(h.title) < 3 or len(h.title) > 200:
                continue

            # 检查层级合理性
            if h.level > prev_level + 1:
                # 层级跳跃，尝试修复
                h.level = prev_level + 1
                if prev_numbering:
                    h.numbering = prev_numbering + [1]
                    h.numbering_str = '.'.join(map(str, h.numbering))

            # 检查编号连续性
            if h.level == prev_level and prev_numbering:
                expected = prev_numbering.copy()
                expected[-1] += 1
                if h.numbering != expected:
                    # 宽松模式：允许跳跃，不做强制修复
                    pass

            repaired.append(h)
            prev_level = h.level
            prev_numbering = h.numbering

        return repaired

    def _calculate_end_positions(self, headings: List[Heading], total_paras: int) -> List[Heading]:
        """计算每个章节的结束位置"""
        for i, h in enumerate(headings):
            if i < len(headings) - 1:
                h.end_para = headings[i + 1].start_para
            else:
                h.end_para = total_paras
        return headings


class SectionExtractor:
    """
    内容提取器 - 层级包含提取

    提取指定章节及其所有子节
    """

    def extract(self, tree: HeadingTree, section_id: str) -> List[Dict]:
        """
        提取指定章节内容（含子节）

        Args:
            tree: 章节树
            section_id: 章节号，如 "2.1"

        Returns:
            List[Dict]: 章节内容列表
        """
        # 查找目标章节
        target = tree.find_section(section_id)
        if not target:
            return []

        # 收集目标章节及其子节
        sections_to_extract = [target]
        sections_to_extract.extend(tree.get_children(target))

        # 提取内容
        results = []
        doc = Document(tree.doc_path)

        for section in sections_to_extract:
            content = self._extract_content(doc, section)
            results.append({
                'level': section.level,
                'title': section.title,
                'numbering': section.numbering,
                'numbering_str': section.numbering_str,
                'content': content,
                'start_para': section.start_para,
                'end_para': section.end_para
            })

        return results

    def _extract_content(self, doc: Document, heading: Heading) -> str:
        """提取章节内容"""
        paragraphs = []

        for i in range(heading.start_para, heading.end_para):
            if i < len(doc.paragraphs):
                text = doc.paragraphs[i].text.strip()
                if text:
                    paragraphs.append(text)

        # 移除标题行本身（如果是第一个段落）
        if paragraphs and heading.title in paragraphs[0]:
            paragraphs = paragraphs[1:]

        return '\n\n'.join(paragraphs)


class MemoryUpdater:
    """
    记忆更新器

    - 生成 Markdown diff
    - 用户审核
    - 备份写入
    """

    def generate_diff(
        self,
        memory_path: Optional[Path],
        section_content: List[Dict],
        section_id: str
    ) -> str:
        """
        生成 Markdown 格式的 diff

        Args:
            memory_path: 现有记忆文件路径（可能不存在）
            section_content: 提取的章节内容
            section_id: 章节号

        Returns:
            str: Markdown diff
        """
        lines = []

        # 标题
        if memory_path and memory_path.exists():
            lines.append(f"### 章节 {section_id} 更新对比")
        else:
            lines.append(f"### 章节 {section_id} 新增内容")

        lines.append("")

        # 章节统计
        total_paras = sum(len(s['content'].split('\n\n')) for s in section_content)
        lines.append(f"**章节**: {section_id}")
        lines.append(f"**包含子节**: {len(section_content)}")
        lines.append(f"**段落数**: {total_paras}")
        lines.append("")

        # 内容预览
        lines.append("**提取内容预览**:")
        lines.append("")

        for section in section_content:
            indent = "  " * (section['level'] - 1)
            lines.append(f"{indent}- **{section['numbering_str']}** {section['title']}")

            # 显示前3个段落
            preview_paras = section['content'].split('\n\n')[:3]
            for para in preview_paras:
                if len(para) > 100:
                    para = para[:100] + "..."
                lines.append(f"{indent}  > {para}")

        lines.append("")

        # 对比表格（如果存在原内容）
        if memory_path and memory_path.exists():
            old_content = memory_path.read_text(encoding='utf-8')

            lines.append("**变更对比**:")
            lines.append("")
            lines.append("| 操作 | 内容 |")
            lines.append("|------|------|")

            # 简单对比：检查章节是否已存在
            section_header = f"## {section_id}"
            if section_header in old_content:
                lines.append(f"| 🟡 修改 | 章节 {section_id} 内容更新 |")
            else:
                lines.append(f"| 🟢 新增 | 章节 {section_id} |")

        return '\n'.join(lines)

    def write_section(
        self,
        memory_path: Path,
        section_content: List[Dict],
        section_id: str
    ):
        """
        写入章节到记忆文件

        Args:
            memory_path: 记忆文件路径
            section_content: 章节内容
            section_id: 章节号
        """
        # 读取现有内容
        if memory_path.exists():
            old_content = memory_path.read_text(encoding='utf-8')
        else:
            old_content = self._create_memory_template()

        # 生成新章节内容
        new_section = self._format_section(section_content, section_id)

        # 替换或追加
        updated_content = self._replace_or_append(old_content, section_id, new_section)

        # 写入文件
        memory_path.write_text(updated_content, encoding='utf-8')

    def _create_memory_template(self) -> str:
        """创建记忆文件模板"""
        return """---
type: thesis-memory
status: 进行中
created: {date}
updated: {date}
---

# 论文记忆文件

## 章节状态总览

| 章节 | 标题 | 状态 |
|------|------|------|

""".format(date=Path().stat().st_ctime if Path().exists() else "")

    def _format_section(self, section_content: List[Dict], section_id: str) -> str:
        """
        格式化章节内容为记忆文件格式（结构化摘要，非全文）

        记忆文件格式：
        - 章节标题
        - 状态标记
        - 核心论点（1-2句）
        - 关键内容（要点列表）
        - 重要引用
        - 字数/段落统计
        """
        lines = []
        import re

        # 主章节标题
        main_section = section_content[0]
        main_title = main_section['title']
        main_numbering = main_section['numbering_str']

        lines.append(f"### {main_numbering} {main_title}")
        lines.append("")

        # 状态标记（根据内容长度判断）
        total_chars = sum(len(s['content']) for s in section_content)
        if total_chars > 3000:
            status = "🟢 已完成"
        elif total_chars > 500:
            status = "🟡 部分完成"
        else:
            status = "🔴 空/占位符"
        lines.append(f"**状态**：{status}")
        lines.append("")

        # 统计信息
        total_paras = sum(len(s['content'].split('\n\n')) for s in section_content)
        lines.append(f"**字数**：{total_chars} 字 | **段落**：{total_paras} 段")
        lines.append("")

        # 处理每个子节
        for section in section_content:
            level = section['level']
            title = section['title']
            content = section['content']
            numbering = section['numbering_str']

            # 跳过主章节本身（已在上面处理）
            if level == main_section['level']:
                continue

            # 子节标题
            lines.append(f"#### {numbering} {title}")
            lines.append("")

            # 内容摘要（如果内容足够）
            if len(content) > 100:
                # 提取核心论点（第一段）
                paras = content.split('\n\n')
                first_para = paras[0][:200] if paras[0] else ""

                if first_para:
                    lines.append(f"**核心论点**：")
                    lines.append(f"- {first_para}{'...' if len(paras[0]) > 200 else ''}")
                    lines.append("")

                # 关键内容（提取要点）
                lines.append(f"**关键内容**：")

                # 查找列表项（以 - 或 ├─ 或 1. 开头）
                key_points = []
                for para in paras[1:8]:  # 限制提取前8个要点
                    para = para.strip()
                    if len(para) < 20:  # 跳过太短的段落
                        continue
                    if len(para) > 150:
                        para = para[:150] + "..."
                    key_points.append(para)
                    if len(key_points) >= 5:  # 最多5个要点
                        break

                if key_points:
                    for point in key_points:
                        lines.append(f"- {point}")
                else:
                    # 如果没有明显要点，提取前3段
                    for para in paras[1:4]:
                        if len(para) > 30:
                            if len(para) > 120:
                                para = para[:120] + "..."
                            lines.append(f"- {para}")

                lines.append("")

                # 重要引用（查找 [数字] 格式的引用）
                citations = re.findall(r'\[(\d+)\]', content)
                if citations:
                    unique_citations = sorted(set(citations), key=int)[:10]  # 最多10个
                    lines.append(f"**关键引用**：[{'], ['.join(unique_citations)}]")
                    lines.append("")
            else:
                # 内容太少，标记为占位符
                lines.append(f"**内容**：{content[:100]}")
                lines.append("")

        return '\n'.join(lines)

        return '\n'.join(lines)

    def _replace_or_append(self, old_content: str, section_id: str, new_section: str) -> str:
        """替换或追加章节"""
        # 查找章节标记
        section_pattern = f"## {section_id}"

        if section_pattern in old_content:
            # 替换现有章节
            # 找到章节开始位置
            start_idx = old_content.find(section_pattern)

            # 找到下一个同级或更高级章节
            next_section_pattern = f"\n## "
            end_idx = old_content.find(next_section_pattern, start_idx + len(section_pattern))

            if end_idx == -1:
                # 这是最后一个章节
                return old_content[:start_idx] + new_section
            else:
                return old_content[:start_idx] + new_section + old_content[end_idx:]
        else:
            # 追加到文件末尾
            return old_content.rstrip() + "\n\n" + new_section

    def generate_full_preview(self, docx_path: str, tree: HeadingTree, all_sections: List[List[Dict]]) -> str:
        """生成全量扫描预览"""
        lines = []

        lines.append("# 论文全量扫描预览")
        lines.append("")
        lines.append(f"**源文件**: {docx_path}")
        lines.append(f"**章节数**: {len(tree.get_level1_sections())}")
        lines.append("")

        # 章节结构树
        lines.append("## 章节结构")
        lines.append("")

        for heading in tree.headings:
            indent = "  " * (heading.level - 1)
            status = "🟢" if heading.level == 1 else "➖"
            lines.append(f"{indent}{status} **{heading.numbering_str}** {heading.title}")

        return '\n'.join(lines)
