#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word Converter
将Markdown文件转换为Word文档，保持Obsidian风格的格式
支持表格、标题、列表、代码块等元素
中文字体使用黑体和宋体
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants


class MarkdownToWord:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()

    def setup_styles(self):
        """设置中文字体和样式"""
        # 设置文档默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 设置标题样式
        for i in range(1, 7):
            try:
                heading_style = self.doc.styles[f'Heading {i}']
                heading_style.font.name = '黑体'
                heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                heading_style.font.bold = True
                heading_style.font.color.rgb = RGBColor(0, 0, 0)

                # 设置标题字号
                if i == 1:
                    heading_style.font.size = Pt(22)
                elif i == 2:
                    heading_style.font.size = Pt(18)
                elif i == 3:
                    heading_style.font.size = Pt(16)
                elif i == 4:
                    heading_style.font.size = Pt(14)
                else:
                    heading_style.font.size = Pt(12)
            except KeyError:
                pass

    def add_hyperlink(self, paragraph, text, url):
        """添加超链接（简化版，避免导入问题）"""
        # 简化处理：直接显示文本和URL
        run = paragraph.add_run(f"{text} ({url})")
        run.font.color.rgb = RGBColor(5, 99, 193)
        run.font.underline = True
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        return run

    def parse_inline_formatting(self, text):
        """解析内联格式（加粗、斜体、行内代码、超链接等）"""
        # 先移除转义字符
        text = text.replace('\\*', '\x00asterisk\x00')
        text = text.replace('\\_', '\x00underscore\x00')
        text = text.replace('\\`', '\x00backtick\x00')

        patterns = []

        # 加粗 **text**
        for match in re.finditer(r'\*\*(.+?)\*\*', text):
            patterns.append((match.start(), match.end(), 'bold', match.group(1)))

        # 斜体 *text* 或 _text_
        for match in re.finditer(r'(?<![\*\w])\*(.+?)\*(?![\*\w])', text):
            if not (match.start() > 0 and text[match.start()-1] == '*'):
                patterns.append((match.start(), match.end(), 'italic', match.group(1)))

        for match in re.finditer(r'_(.+?)_', text):
            patterns.append((match.start(), match.end(), 'italic', match.group(1)))

        # 行内代码 `code`
        for match in re.finditer(r'`([^`]+)`', text):
            patterns.append((match.start(), match.end(), 'code', match.group(1)))

        # 删除线 ~~text~~
        for match in re.finditer(r'~~(.+?)~~', text):
            patterns.append((match.start(), match.end(), 'strike', match.group(1)))

        # 按位置排序
        patterns.sort(key=lambda x: x[0])

        # 移除重叠
        result = []
        last_end = 0
        for start, end, fmt, content in patterns:
            if start >= last_end:
                result.append((start, end, fmt, content))
                last_end = end

        return result

    def add_formatted_text(self, paragraph, text):
        """添加带格式的文本到段落"""
        # 解析内联格式
        formats = self.parse_inline_formatting(text)

        if not formats:
            # 没有特殊格式，直接添加
            run = paragraph.add_run(text)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            return

        # 按格式分段添加
        last_end = 0
        for start, end, fmt, content in formats:
            # 添加前面的普通文本
            if start > last_end:
                normal_text = text[last_end:start]
                normal_text = normal_text.replace('\x00asterisk\x00', '*')
                normal_text = normal_text.replace('\x00underscore\x00', '_')
                normal_text = normal_text.replace('\x00backtick\x00', '`')
                run = paragraph.add_run(normal_text)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 添加格式化文本
            content = content.replace('\x00asterisk\x00', '*')
            content = content.replace('\x00underscore\x00', '_')
            content = content.replace('\x00backtick\x00', '`')
            run = paragraph.add_run(content)

            if fmt == 'bold':
                run.bold = True
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            elif fmt == 'italic':
                run.italic = True
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            elif fmt == 'code':
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(199, 37, 78)
                # 添加浅灰背景
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F5F5F5')
                run._element.rPr.append(shading)
            elif fmt == 'strike':
                run.font.strike = True
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            last_end = end

        # 添加剩余文本
        if last_end < len(text):
            remaining = text[last_end:]
            remaining = remaining.replace('\x00asterisk\x00', '*')
            remaining = remaining.replace('\x00underscore\x00', '_')
            remaining = remaining.replace('\x00backtick\x00', '`')
            run = paragraph.add_run(remaining)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def process_table(self, lines, start_idx):
        """处理Markdown表格"""
        table_lines = []
        i = start_idx

        # 收集表格行
        while i < len(lines) and '|' in lines[i]:
            table_lines.append(lines[i])
            i += 1

        if len(table_lines) < 2:
            return i

        # 解析表头
        header_line = table_lines[0]
        headers = [cell.strip() for cell in header_line.split('|')[1:-1]]

        # 创建表格
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # 设置表头
        hdr_cells = table.rows[0].cells
        for j, header in enumerate(headers):
            hdr_cells[j].text = header
            # 设置表头字体
            for paragraph in hdr_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = Pt(11)

        # 跳过分隔线（第二行）
        data_start = 2

        # 处理数据行
        for line in table_lines[data_start:]:
            if line.strip():
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                row_cells = table.add_row().cells
                for j, cell in enumerate(cells):
                    if j < len(row_cells):
                        row_cells[j].text = cell
                        # 设置单元格字体
                        for paragraph in row_cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.font.name = '宋体'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                run.font.size = Pt(10.5)

        return i

    def process_code_block(self, lines, start_idx):
        """处理代码块"""
        first_line = lines[start_idx]
        lang_match = re.match(r'```(\w*)', first_line)
        language = lang_match.group(1) if lang_match else ''

        code_lines = []
        i = start_idx + 1

        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1

        # 添加代码块段落
        code_text = '\n'.join(code_lines)

        # 添加代码块容器（带边框的段落）
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(51, 51, 51)

        # 添加代码块边框（通过底纹模拟）
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F8F8F8')
        p._p.get_or_add_pPr().append(shading)

        return i + 1 if i < len(lines) else i

    def convert(self, md_content):
        """转换Markdown内容到Word"""
        lines = md_content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # 跳过空行
            if not line.strip():
                i += 1
                continue

            # 代码块
            if line.startswith('```'):
                i = self.process_code_block(lines, i)
                continue

            # 表格
            if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
                i = self.process_table(lines, i)
                continue

            # 六级标题 ######
            if line.startswith('###### '):
                p = self.doc.add_heading(line[7:], level=6)
                self.set_chinese_font(p, '黑体', Pt(12))
                i += 1
                continue

            # 五级标题 #####
            if line.startswith('##### '):
                p = self.doc.add_heading(line[6:], level=5)
                self.set_chinese_font(p, '黑体', Pt(12))
                i += 1
                continue

            # 四级标题 ####
            if line.startswith('#### '):
                p = self.doc.add_heading(line[5:], level=4)
                self.set_chinese_font(p, '黑体', Pt(14))
                i += 1
                continue

            # 三级标题 ###
            if line.startswith('### '):
                p = self.doc.add_heading(line[4:], level=3)
                self.set_chinese_font(p, '黑体', Pt(16))
                i += 1
                continue

            # 二级标题 ##
            if line.startswith('## '):
                p = self.doc.add_heading(line[3:], level=2)
                self.set_chinese_font(p, '黑体', Pt(18))
                i += 1
                continue

            # 一级标题 #
            if line.startswith('# ') and not line.startswith('## '):
                p = self.doc.add_heading(line[2:], level=1)
                self.set_chinese_font(p, '黑体', Pt(22))
                i += 1
                continue

            # 引用块 >
            if line.startswith('>'):
                quote_text = line[1:].strip()
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                self.add_formatted_text(p, quote_text)
                i += 1
                continue

            # 无序列表 -, *, +
            list_match = re.match(r'^[\s]*([\-\*\+])\s+(.+)', line)
            if list_match:
                indent_level = len(line) - len(line.lstrip())
                content = list_match.group(2)
                p = self.doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25 * (indent_level // 2 + 1))
                self.add_formatted_text(p, content)
                i += 1
                continue

            # 有序列表 1., 2., etc.
            ordered_match = re.match(r'^[\s]*(\d+)\.\s+(.+)', line)
            if ordered_match:
                indent_level = len(line) - len(line.lstrip())
                content = ordered_match.group(2)
                p = self.doc.add_paragraph(style='List Number')
                p.paragraph_format.left_indent = Inches(0.25 * (indent_level // 2 + 1))
                self.add_formatted_text(p, content)
                i += 1
                continue

            # 水平分隔线 ---, ***, ___
            if re.match(r'^[\s]*([\-\*_][\s]*){3,}[\s]*$', line):
                self.doc.add_paragraph('─' * 50)
                i += 1
                continue

            # 普通段落
            p = self.doc.add_paragraph()
            self.add_formatted_text(p, line)
            i += 1

    def set_chinese_font(self, paragraph, font_name, font_size):
        """设置段落的中文字体"""
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if font_size:
                run.font.size = font_size

    def save(self, output_path):
        """保存Word文档"""
        self.doc.save(output_path)
        print(f"文档已保存至: {output_path}")


def main():
    if len(sys.argv) < 2:
        print("用法: python md_to_word.py <输入.md> [输出.docx]")
        print("示例: python md_to_word.py readme.md readme.docx")
        sys.exit(1)

    input_file = sys.argv[1]

    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        output_file = str(Path(input_file).with_suffix('.docx'))

    if not Path(input_file).exists():
        print(f"错误: 找不到文件 '{input_file}'")
        sys.exit(1)

    print(f"正在转换: {input_file} -> {output_file}")

    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        converter = MarkdownToWord()
        converter.convert(md_content)
        converter.save(output_file)
        print("转换完成!")

    except Exception as e:
        print(f"转换出错: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
